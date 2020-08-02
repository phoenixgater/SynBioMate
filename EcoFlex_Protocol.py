# Import libraries
from docx import Document

# Import scripts
import GUI
import MoClo

# Document for writing protocol with docx (This is NOT a pySBOL document)
document = Document()

# Dictionaries for well allocations
level_1_384PP = {}
level_1_LDV = {}
level_1_6RES = {}
level_1_output = {}
level_2_384PP = {}
level_2_LDV = {}
level_2_6RES = {}
level_2_output = {}
# Plate lists
well_letters_384 = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P"]
well_numbers_384 = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12", "13", "14", "15", "16", "17", "18",
                    "19", "20", "21", "22", "23", "24"]


def get_variables():
    # Protocol name
    global protocol_name
    protocol_name = (str(GUI.protocol_name_entry.get())).replace(" ", "_")

    # Imported user input parameters
    global transcription_unit_quantity
    transcription_unit_quantity = GUI.transcription_unit_quantity_combo.get()
    global signal_peptide_choice
    signal_peptide_choice = GUI.include_signal_combo.get()


def create_selected_reaction_list():
    global selected_reaction_ratios_level_1
    selected_reaction_ratios_level_1 = []
    if GUI.level_1_ratio_1_1.get() == 1:
        selected_reaction_ratios_level_1.append(" reaction 1:1")
    if GUI.level_1_ratio_1_2.get() == 1:
        selected_reaction_ratios_level_1.append(" reaction 1:2")
    if GUI.level_1_ratio_2_1.get() == 1:
        selected_reaction_ratios_level_1.append(" reaction 2:1")
    if not selected_reaction_ratios_level_1:
        selected_reaction_ratios_level_1.append(" reaction 2:1")

    global selected_reaction_ratios_level_2
    selected_reaction_ratios_level_2 = []
    if GUI.level_2_ratio_1_1.get() == 1:
        selected_reaction_ratios_level_2.append(" reaction 1:1")
    if GUI.level_2_ratio_1_2.get() == 1:
        selected_reaction_ratios_level_2.append(" reaction 1:2")
    if GUI.level_2_ratio_2_1.get() == 1:
        selected_reaction_ratios_level_2.append(" reaction 2:1")
    if not selected_reaction_ratios_level_2:
        selected_reaction_ratios_level_2.append(" reaction 2:1")


# Assign wells to genetic parts and reagents for 384 plates, specifies volumes
def assign_well_384(plate_dictionary, item, volume):
    if not plate_dictionary.keys():
        plate_dictionary["A1"] = [item, volume]
    else:
        taken_wells = list(plate_dictionary.keys())
        if len(taken_wells[-1]) == 2:
            pos = well_numbers_384.index(taken_wells[-1][1])
            new_number = well_numbers_384[pos + 1]
            plate_dictionary[taken_wells[-1][0] + new_number] = [item, volume]
        else:
            if taken_wells[-1][-2:] == "24":
                pos = well_letters_384.index(taken_wells[-1][0])
                new_letter = well_letters_384[pos + 1]
                plate_dictionary[new_letter + "1"] = [item, volume]
            else:
                pos = well_numbers_384.index(taken_wells[-1][-2:])
                new_number = well_numbers_384[pos + 1]
                plate_dictionary[taken_wells[-1][0] + new_number] = [item, volume]


# Assign wells to reagents for 6RES plates, specifies volumes
def assign_well_6res(plate_dictionary, item, volume):
    if not plate_dictionary.keys():
        plate_dictionary["A1"] = [item, volume]
    else:
        taken_wells = list(plate_dictionary.keys())
        if taken_wells[-1] == "A1":
            plate_dictionary["A2"] = [item, volume]
        elif taken_wells[-1] == "A2":
            plate_dictionary["A3"] = [item, volume]
        elif taken_wells[-1] == "A3":
            plate_dictionary["B1"] = [item, volume]
        elif taken_wells[-1] == "B1":
            plate_dictionary["B2"] = [item, volume]
        elif taken_wells[-1] == "B2":
            plate_dictionary["B3"] = [item, volume]


# Returns single transfer volumes that correspond to the selected user-input parameters "1_1" denotes the reaction ratio
def transfer_volume(item):
    if item == "level_1_water_2_1":
        if GUI.include_signal_combo.get() == "No":
            return 1875
        else:
            return 1375
    elif item == "level_1_water_1_2":
        if GUI.include_signal_combo.get() == "No":
            return 2625
        else:
            return 2375
    elif item == "level_1_water_1_1":
        if GUI.include_signal_combo.get() == "No":
            return 2875
        else:
            return 2625

    elif item == "level_1_part_2_1":
        return 500
    elif item == "level_1_part_1_2":
        return 250
    elif item == "level_1_part_1_1":
        return 250

    elif item == "level_1_backbone_2_1":
        return 250
    elif item == "level_1_backbone_1_2":
        return 500
    elif item == "level_1_backbone_1_1":
        return 250

    elif item == "level_2_water_2_1":
        if GUI.transcription_unit_quantity_combo.get() == "2":
            return 2875
        elif GUI.transcription_unit_quantity_combo.get() == "3":
            return 2375
        elif GUI.transcription_unit_quantity_combo.get() == "4":
            return 1875
        elif GUI.transcription_unit_quantity_combo.get() == "5":
            return 1375
    elif item == "level_2_water_1_2":
        if GUI.transcription_unit_quantity_combo.get() == "2":
            return 3125
        elif GUI.transcription_unit_quantity_combo.get() == "3":
            return 2875
        elif GUI.transcription_unit_quantity_combo.get() == "4":
            return 2625
        elif GUI.transcription_unit_quantity_combo.get() == "5":
            return 2375
    elif item == "level_2_water_1_1":
        if GUI.transcription_unit_quantity_combo.get() == "2":
            return 3375
        elif GUI.transcription_unit_quantity_combo.get() == "3":
            return 3125
        elif GUI.transcription_unit_quantity_combo.get() == "4":
            return 2875
        elif GUI.transcription_unit_quantity_combo.get() == "5":
            return 2625

    elif item == "level_2_tu_2_1":
        return 500
    elif item == "level_2_tu_1_2":
        return 250
    elif item == "level_2_tu_1_1":
        return 250

    elif item == "level_2_backbone_2_1":
        return 250
    elif item == "level_2_backbone_1_2":
        return 500
    elif item == "level_2_backbone_1_1":
        return 250


# Calculate part quantity
def calculate_part_quantity():
    global part_quantity
    if GUI.include_signal_combo.get() == "Yes":
        part_quantity = (len(MoClo.promoter_identities) + len(MoClo.rbs_identities) + len(MoClo.signal_identities)
                         + len(MoClo.cds_identities) + len(MoClo.terminator_identities))
    elif GUI.include_signal_combo.get() == "No":
        part_quantity = (len(MoClo.promoter_identities) + len(MoClo.rbs_identities)
                         + len(MoClo.cds_identities) + len(MoClo.terminator_identities))


# Calculate quantity of level 1 transcription units
def calculate_level_1_quantity():
    global level_1_tu_quantity
    if int(GUI.transcription_unit_quantity_combo.get()) == 2:
        level_1_tu_quantity = len(MoClo.transcription_unit_1_names) + len(MoClo.transcription_unit_2_names)
    if int(GUI.transcription_unit_quantity_combo.get()) == 3:
        level_1_tu_quantity = (len(MoClo.transcription_unit_1_names) + len(MoClo.transcription_unit_2_names) +
                               len(MoClo.transcription_unit_3_names))
    if int(GUI.transcription_unit_quantity_combo.get()) == 4:
        level_1_tu_quantity = (len(MoClo.transcription_unit_1_names) + len(MoClo.transcription_unit_2_names) +
                               len(MoClo.transcription_unit_3_names) + len(MoClo.transcription_unit_4_names))
    if int(GUI.transcription_unit_quantity_combo.get()) == 5:
        level_1_tu_quantity = (len(MoClo.transcription_unit_1_names) + len(MoClo.transcription_unit_2_names) +
                               len(MoClo.transcription_unit_3_names) + len(MoClo.transcription_unit_4_names) +
                               len(MoClo.transcription_unit_5_names))


# Title and introduction of document
def title_introduction():
    if GUI.assembly_method_combo.get() == "Automatic":
        document.add_heading("Automated MoClo assembly protocol", 0)
        document.add_paragraph("Hello! This document contains a protocol for the assembly of genetic parts using" +
                               " MoClo assembly with an automated liquid handler. This document was produced by the" +
                               " software 'SynBioMate' (https://github.com/phoenixgater/SynBioMate)")
    elif GUI.assembly_method_combo.get() == "Manual":
        document.add_heading("Manual MoClo assembly protocol", 0)
        document.add_paragraph("Hello! This document contains a protocol for the manual assembly of genetic parts" +
                               " using MoClo assembly. This document was produced by the" +
                               " software 'SynBioMate' (https://github.com/phoenixgater/SynBioMate)")
    document.add_paragraph("The protocols and fusion sites in this protocol are designed to be compatible with" +
                           " the EcoFlex MoClo kit, " "available from: http://www.addgene.org/kits/freemont-ecoflex"
                           "-moclo/")

    part_notes = document.add_paragraph("")
    part_notes.add_run("Notes on this document and its contents:").bold = True
    part_notes.add_run("\n" + "-For restriction sites detected in open reading frames (coding regions, CDSs, ORFs)," +
                       " a codon in this region has been swapped to ensure that the part is MoClo compatible." +
                       " This codon swap is specified in this documents appendix.")
    part_notes.add_run("\n" + "\n" + "-If an ATG start codon could not" +
                       " be found at the start of an ORF/CDS part, this has been added, and noted in the appendix as" +
                       " well.")
    part_notes.add_run(" It is, however, highly advised that the use of these parts be reconsidered").bold = True
    part_notes.add_run("\n" + "\n" + "-This software is unable to suggest base substitutions for excluded restriction "
                                     "sites" +
                       " detected in non-coding genetic parts (e.g Promoters, RBSs, etc), but the presence" +
                       " of these restriction sites will be noted in this documents appendix.")
    part_notes.add_run(" It is highly advised that the use of these parts be reconsidered").bold = True

    input_notes = document.add_paragraph("Please note that the contents of this document will only be valid for " +
                                         "the user-input parameters that were given at the time of this documents " +
                                         "creation. These were:")
    input_notes.add_run("\n" + "Signal peptide included:").bold = True
    input_notes.add_run(" " + GUI.include_signal_combo.get())
    input_notes.add_run("\n" + "Chassis system:").bold = True
    input_notes.add_run(" " + GUI.chassis_selection_combo.get())
    input_notes.add_run("\n" + "Transcription unit quantity:").bold = True
    input_notes.add_run(" " + GUI.transcription_unit_quantity_combo.get())
    input_notes.add_run("\n" + "Assembly method:").bold = True
    input_notes.add_run(" " + GUI.assembly_method_combo.get())
    if GUI.assembly_method_combo.get() == "Automatic":
        input_notes.add_run("\n" + "Liquid handler:").bold = True
        input_notes.add_run(" " + GUI.liquid_handler_selection_combo.get())
    input_notes.add_run("\n" + "Include codon swap for cds parts?:").bold = True
    input_notes.add_run(" " + GUI.include_codon_swap_combo.get())
    input_notes.add_run("\n" + "Add fusion sites to parts?:").bold = True
    input_notes.add_run(" " + GUI.include_fusion_site_combo.get())
    input_notes.add_run("\n" + "Selected reaction ratios:").bold = True
    input_notes.add_run("\n" + "Level 1: ")
    for ratio in selected_reaction_ratios_level_1:
        input_notes.add_run(ratio + ", ")
    input_notes.add_run("\n" + "Level 2: ")
    for ratio in selected_reaction_ratios_level_2:
        input_notes.add_run(ratio + ", ")
    document.add_page_break()


# Write manual protocol to word document
def create_manual_protocol():
    document.add_heading("Protocol", 1)
    document.add_heading("Creating level 0 library and cloning into level 0 plasmid backbones", 2)
    biopart_prep = document.add_paragraph("")
    biopart_prep.add_run("a) For each of the BioParts in this documents appendix (all parts EXCLUDING CDS/ORF): ")
    biopart_prep.add_run("\n" + "1. Design and create/purchase forward and reverse primers for the BioPart")
    biopart_prep.add_run("\n" + "2. Anneal forward and reverse primers (20 uM of each) in 1 x T4 DNA ligase buffer " +
                         "at 90°C for 1 min, followed by cooling on ice")
    biopart_prep.add_run("\n" + "3. Phosphorylate with T4 PNK for 1 hour at 37°C")
    biopart_prep.add_run("\n" + "4. Digest 40 ng/uL-1 (24nM) of pBP-lacZ (available from "
                                "http://www.addgene.org/72948/)  with 1 unit of NdeI and SphI " +
                         "in CutSmart buffer for 1 hour at 37°C, heat inactivate at 80°C for 10 min")
    biopart_prep.add_run("\n" + "5. Add 2 uL (80 ng) of NdeI-SphI digested pBP-lacZ plasmid with 2 uL of 200 nM "
                         + "annealed primers in 2 x Rapid ligation buffer (Promega) and 1-3 units of T4 ligase " +
                         "(Promega)")
    biopart_prep.add_run("\n" + "6. Incubate at room-temperature (19-21 °C) for 30-60 min")
    biopart_prep.add_run("\n" + "7. Transform 5 uL into 50 uL of DH10a competent cells")
    biopart_prep.add_run("\n" + "8. Grow on 35 ug mL-1 chloramphenicol LB plates with 0.1 mM IPTG and 40 ug " +
                         "mL-1 X-Gal (Sigma) for 24 hrs")
    biopart_prep.add_run("\n" + "9. Place in cold room overnight for clear visualisation of negative blue colonies")
    biopart_prep.add_run("\n" + "10. Pick 1-2 white colonies for plasmid preparation and sequencing")

    biopart_prep.add_run("\n" + "\n" + "b) For each of the CDS/ORF parts in this documents appendix: ")
    biopart_prep.add_run("\n" + "1. Design and create/purchase forward and reverse primers for the CDS part")
    biopart_prep.add_run("\n" + "Anneal forward and reverse primers (20 uM of each) in 1 x T4 DNA ligase buffer at " +
                         "90°C for 1 min, followed by cooling on ice")
    biopart_prep.add_run("\n" + "3. Phosphorylate with T4 PNK for 1 hour at 37°C")
    biopart_prep.add_run("\n" + "4. Digest 40 ng/uL-1 (24nM) of pBP-ORF (available from "
                                "http://www.addgene.org/72949/) with 1 unit of NdeI and BamHI " +
                         "in CutSmart buffer for 1 hour at 37°C, heat inactivate at 80°C for 10 min")
    biopart_prep.add_run("\n" + "5. Add 2 uL (80 ng) of NdeI-BamHI digested pBP-ORF plasmid with 2 uL of 200 nM " +
                         "annealed primers in 2 x Rapid ligation buffer (Promega) and 1-3 units of T4 ligase (Promega)")
    biopart_prep.add_run("\n" + "6. Incubate at room-temperature (19-21 °C) for 30-60 min")
    biopart_prep.add_run("\n" + "7. Transform 5 uL into 50 uL of DH10a competent cells")
    biopart_prep.add_run("\n" + "8. Grow on 35 ug mL-1 chloramphenicol LB plates with 0.1 mM IPTG and 40 ug" +
                         " mL-1 X-Gal (Sigma) for 24 hrs")
    biopart_prep.add_run("\n" + "9. Place in cold room overnight for clear visualisation of negative blue colonies")
    biopart_prep.add_run("\n" + "10. Pick 1-2 white colonies for plasmid preparation and sequencing")

    document.add_heading("Creating level 1 transcription units (TUs)", 2)
    tu_prep = document.add_paragraph("")
    tu_prep.add_run("For each level 1 transcription unit variant in this documents appendix:")
    tu_prep.add_run("\n" + "Create a mixture consisting of:" + "\n" + "-100ng of each included part" +
                    "\n" + "-50ng of the level 1 destination plasmid backbone (Backbones will differ for each TU," +
                    " the backbone to be used for each respective TU is specified in the appendix)" + "\n" +
                    "-20 units BsaI-HF (NEB)" + "\n" + "-1-3 units T4 DNA ligase (Promega)")
    tu_prep.add_run("\n" + "\n" + "b) Run a PCR protocol for this mixture, consisting of:" + "\n" + "15-30 cycles of:" +
                    "\n" + "-5 minutes at 37°C" + "\n" + "-10 minutes at 16°C" + "\n" + "Followed by:" + "\n" +
                    "-5 minutes at 50°C" + "\n" + "-5 minutes at 80°C")
    tu_prep.add_run("\n" + "c) Transform 5 uL of the mixture into 50uL of chemically competent" +
                    "Escherichia coli (E. coli) dH10a by heat shock transformation")

    document.add_heading("Creating level 2 transcription units (TUs)", 2)
    tu2_prep = document.add_paragraph("")
    tu2_prep.add_run("For each level 2 transcription unit variant in this document:")
    tu2_prep.add_run("\n" + "a) Create a mixture consisting of:")
    tu2_prep.add_run("\n" + "-100ng of each included level 1 transcription unit" + "\n" +
                     "-50ng of level 2 destination vector" + "\n" + "-10 units BsmBI (NEB)" + "\n" +
                     "-1-3 units T4 DNA ligase (Promega)")
    tu2_prep.add_run("\n" + "\n" + "b) Incubate mixture at 37°C for 16 hours")
    tu2_prep.add_run("\n" + "c) Incubate at 80°C for 5 minutes")
    tu2_prep.add_run("\n" + "d) Transform 5 uL of the mixture into 50uL of chemically competent E. coli dH10a " +
                     "by heat shock transformation")


# Write automatic protocol to word document, create plate dictionaries
def create_automatic_protocol():
    document.add_heading("Protocol", 1)
    document.add_heading("Creating level 0 library and cloning into level 0 plasmid backbones", 2)
    biopart_prep = document.add_paragraph("")
    biopart_prep.add_run("a) For each of the BioParts in this documents appendix (all parts EXCLUDING CDS/ORF): ")
    biopart_prep.add_run("\n" + "1. Design and create/purchase forward and reverse primers for the BioPart")
    biopart_prep.add_run(
        "\n" + "2. Anneal forward and reverse primers (20 uM of each) in 1 x T4 DNA ligase buffer " +
        "at 90°C for 1 min, followed by cooling on ice")
    biopart_prep.add_run("\n" + "3. Phosphorylate with T4 PNK for 1 hour at 37°C")
    biopart_prep.add_run("\n" + "4. Digest 40 ng/uL-1 (24nM) of pBP-lacZ (available from "
                                "http://www.addgene.org/72948/) with 1 unit of NdeI and SphI " +
                         "in CutSmart buffer for 1 hour at 37°C, heat inactivate at 80°C for 10 min")
    biopart_prep.add_run("\n" + "5. Add 2 uL (80 ng) of NdeI-SphI digested pBP-lacZ plasmid with 2 uL of 200 nM "
                         + "annealed primers in 2 x Rapid ligation buffer (Promega) and 1-3 units of T4 ligase " +
                         "(Promega)")
    biopart_prep.add_run("\n" + "6. Incubate at room-temperature (19-21 °C) for 30-60 min")
    biopart_prep.add_run("\n" + "7. Transform 5 uL into 50 uL of DH10a competent cells")
    biopart_prep.add_run("\n" + "8. Grow on 35 ug mL-1 chloramphenicol LB plates with 0.1 mM IPTG and 40 ug " +
                         "mL-1 X-Gal (Sigma) for 24 hrs")
    biopart_prep.add_run("\n" + "9. Place in cold room overnight for clear visualisation of negative blue colonies")
    biopart_prep.add_run("\n" + "10. Pick 1-2 white colonies for plasmid preparation and sequencing")

    biopart_prep.add_run("\n" + "\n" + "b) For each of the CDS/ORF parts in this documents appendix: ")
    biopart_prep.add_run("\n" + "1. Design and create/purchase forward and reverse primers for the CDS part")
    biopart_prep.add_run(
        "\n" + "Anneal forward and reverse primers (20 uM of each) in 1 x T4 DNA ligase buffer at " +
        "90°C for 1 min, followed by cooling on ice")
    biopart_prep.add_run("\n" + "3. Phosphorylate with T4 PNK for 1 hour at 37°C")
    biopart_prep.add_run("\n" + "4. Digest 40 ng/uL-1 (24nM) of pBP-ORF (available from "
                                "http://www.addgene.org/72949/)  with 1 unit of NdeI and BamHI " +
                         "in CutSmart buffer for 1 hour at 37°C, heat inactivate at 80°C for 10 min")
    biopart_prep.add_run("\n" + "5. Add 2 uL (80 ng) of NdeI-BamHI digested pBP-ORF plasmid with 2 uL of 200 nM " +
                         "annealed primers in 2 x Rapid ligation buffer (Promega) and 1-3 units of T4 ligase (Promega)")
    biopart_prep.add_run("\n" + "6. Incubate at room-temperature (19-21 °C) for 30-60 min")
    biopart_prep.add_run("\n" + "7. Transform 5 uL into 50 uL of DH10a competent cells")
    biopart_prep.add_run("\n" + "8. Grow on 35 ug mL-1 chloramphenicol LB plates with 0.1 mM IPTG and 40 ug" +
                         " mL-1 X-Gal (Sigma) for 24 hrs")
    biopart_prep.add_run("\n" + "9. Place in cold room overnight for clear visualisation of negative blue colonies")
    biopart_prep.add_run("\n" + "10. Pick 1-2 white colonies for plasmid preparation and sequencing")

    document.add_page_break()

    document.add_heading("Creating level 1 transcription units (TUs)", 2)
    tu_1_prep_1_intro = document.add_paragraph("")
    tu_1_prep_1_intro.add_run("Plates required:").bold = True
    tu_1_prep_1_intro.add_run("\n" + "-2x Echo® Qualified 384-Well Polypropylene Source Microplate (384PP) (One of "
                                     "these is used as an output plate and must remain empty)")
    tu_1_prep_1_intro.add_run("\n" + "-Echo® Qualified 384-Well COC Source Microplate, Low Dead Volume (384LDV)")
    tu_1_prep_1_intro.add_run("\n" + "-Echo® Qualified Reservoir (6RES)" + "\n")

    tu_1_prep_1 = document.add_paragraph("")
    tu_1_prep_1.add_run("a) Add genetic parts and level 1 transcription unit backbones "
                        "into their corresponding wells in" +
                        " the Echo® Qualified 384-Well Polypropylene Source Microplate (384PP) as" +
                        " specified in the table below")

    ########################## level 1 384 source plate for parts and plasmid backbones ###########################
    level_1_protocol_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_1_protocol_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "Genetic part"
    row_1_cells[2].text = "Quantity (nl)"
    global selected_reaction_ratios_level_1

    # Assigning wells and volumes for parts (level 1)
    part_quantities = MoClo.part_quantities
    dead_volume = 15000
    for key in part_quantities.keys():
        volume_fulfilled = False
        previous_fulfilment = 0
        while not volume_fulfilled:
            required_transfer_volume = 0
            for reaction in selected_reaction_ratios_level_1:
                if reaction == " reaction 1:1":
                    required_transfer_volume += transfer_volume("level_1_part_1_1") * part_quantities[key]
                elif reaction == " reaction 1:2":
                    required_transfer_volume += transfer_volume("level_1_part_1_2") * part_quantities[key]
                elif reaction == " reaction 2:1":
                    required_transfer_volume += transfer_volume("level_1_part_2_1") * part_quantities[key]
            row_cells = level_1_protocol_table.add_row().cells
            volume_required = required_transfer_volume - previous_fulfilment
            row_cells[1].text = key
            if dead_volume + volume_required > 65000:
                volume = 65000
                assign_well_384(level_1_384PP, key, volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                previous_fulfilment += 50000
                continue
            else:
                volume = dead_volume + volume_required
                assign_well_384(level_1_384PP, key, volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                volume_fulfilled = True

    # Assigning wells and volumes for level 1 plasmid backbone A (level 1)
    if int(GUI.transcription_unit_quantity_combo.get()) > 1:
        volume_fulfilled = False
        previous_fulfilment = 0
        while not volume_fulfilled:
            required_transfer_volume = 0
            for reaction in selected_reaction_ratios_level_1:
                if reaction == " reaction 1:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_1") * len(
                        MoClo.transcription_unit_1_names)
                elif reaction == " reaction 1:2":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_2") * len(
                        MoClo.transcription_unit_1_names)
                elif reaction == " reaction 2:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_2_1") * len(
                        MoClo.transcription_unit_1_names)
            row_cells = level_1_protocol_table.add_row().cells
            volume_required = required_transfer_volume - previous_fulfilment
            row_cells[1].text = "pTU1-A-RFP"
            if dead_volume + volume_required > 65000:
                volume = 65000
                assign_well_384(level_1_384PP, "pTU1-A-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                previous_fulfilment += 50000
                continue
            else:
                volume = dead_volume + volume_required
                assign_well_384(level_1_384PP, "pTU1-A-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                volume_fulfilled = True

        # Assigning wells and volumes for level 1 plasmid backbone B (level 1)
        volume_fulfilled = False
        previous_fulfilment = 0
        while not volume_fulfilled:
            required_transfer_volume = 0
            for reaction in selected_reaction_ratios_level_1:
                if reaction == " reaction 1:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_1") * len(
                        MoClo.transcription_unit_2_names)
                elif reaction == " reaction 1:2":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_2") * len(
                        MoClo.transcription_unit_2_names)
                elif reaction == " reaction 2:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_2_1") * len(
                        MoClo.transcription_unit_2_names)
            row_cells = level_1_protocol_table.add_row().cells
            volume_required = required_transfer_volume - previous_fulfilment
            row_cells[1].text = "pTU1-B-RFP"
            if dead_volume + volume_required > 65000:
                volume = 65000
                assign_well_384(level_1_384PP, "pTU1-B-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                previous_fulfilment += 50000
                continue
            else:
                volume = dead_volume + volume_required
                assign_well_384(level_1_384PP, "pTU1-B-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                volume_fulfilled = True

    # Assigning wells and volumes for level 1 plasmid backbone C (level 1)
    if int(GUI.transcription_unit_quantity_combo.get()) > 2:
        volume_fulfilled = False
        previous_fulfilment = 0
        while not volume_fulfilled:
            required_transfer_volume = 0
            for reaction in selected_reaction_ratios_level_1:
                if reaction == " reaction 1:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_1") * len(
                        MoClo.transcription_unit_3_names)
                elif reaction == " reaction 1:2":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_2") * len(
                        MoClo.transcription_unit_3_names)
                elif reaction == " reaction 2:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_2_1") * len(
                        MoClo.transcription_unit_3_names)
            row_cells = level_1_protocol_table.add_row().cells
            volume_required = required_transfer_volume - previous_fulfilment
            row_cells[1].text = "pTU1-C-RFP"
            if dead_volume + volume_required > 65000:
                volume = 65000
                assign_well_384(level_1_384PP, "pTU1-C-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                previous_fulfilment += 50000
                continue
            else:
                volume = dead_volume + volume_required
                assign_well_384(level_1_384PP, "pTU1-C-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                volume_fulfilled = True

    # Assigning wells and volumes for level 1 plasmid backbone D or D1 (level 1)
    if int(GUI.transcription_unit_quantity_combo.get()) > 3:
        if int(GUI.transcription_unit_quantity_combo.get()) == 4:
            volume_fulfilled = False
            previous_fulfilment = 0
            while not volume_fulfilled:
                required_transfer_volume = 0
                for reaction in selected_reaction_ratios_level_1:
                    if reaction == " reaction 1:1":
                        required_transfer_volume += transfer_volume("level_1_backbone_1_1") * len(
                            MoClo.transcription_unit_4_names)
                    elif reaction == " reaction 1:2":
                        required_transfer_volume += transfer_volume("level_1_backbone_1_2") * len(
                            MoClo.transcription_unit_4_names)
                    elif reaction == " reaction 2:1":
                        required_transfer_volume += transfer_volume("level_1_backbone_2_1") * len(
                            MoClo.transcription_unit_4_names)
                row_cells = level_1_protocol_table.add_row().cells
                volume_required = required_transfer_volume - previous_fulfilment
                row_cells[1].text = "pTU1-D-RFP"
                if dead_volume + volume_required > 65000:
                    volume = 65000
                    assign_well_384(level_1_384PP, "pTU1-D-RFP", volume)
                    row_cells[0].text = list(level_1_384PP.keys())[-1]
                    row_cells[2].text = str(volume)
                    previous_fulfilment += 50000
                    continue
                else:
                    volume = dead_volume + volume_required
                    assign_well_384(level_1_384PP, "pTU1-D-RFP", volume)
                    row_cells[0].text = list(level_1_384PP.keys())[-1]
                    row_cells[2].text = str(volume)
                    volume_fulfilled = True

        elif int(GUI.transcription_unit_quantity_combo.get()) == 5:
            volume_fulfilled = False
            previous_fulfilment = 0
            while not volume_fulfilled:
                required_transfer_volume = 0
                for reaction in selected_reaction_ratios_level_1:
                    if reaction == " reaction 1:1":
                        required_transfer_volume += transfer_volume("level_1_backbone_1_1") * len(
                            MoClo.transcription_unit_4_names)
                    elif reaction == " reaction 1:2":
                        required_transfer_volume += transfer_volume("level_1_backbone_1_2") * len(
                            MoClo.transcription_unit_4_names)
                    elif reaction == " reaction 2:1":
                        required_transfer_volume += transfer_volume("level_1_backbone_2_1") * len(
                            MoClo.transcription_unit_4_names)
                row_cells = level_1_protocol_table.add_row().cells
                volume_required = required_transfer_volume - previous_fulfilment
                row_cells[1].text = "pTU1-D1-RFP"
                if dead_volume + volume_required > 65000:
                    volume = 65000
                    assign_well_384(level_1_384PP, "pTU1-D1-RFP", volume)
                    row_cells[0].text = list(level_1_384PP.keys())[-1]
                    row_cells[2].text = str(volume)
                    previous_fulfilment += 50000
                    continue
                else:
                    volume = dead_volume + volume_required
                    assign_well_384(level_1_384PP, "pTU1-D1-RFP", volume)
                    row_cells[0].text = list(level_1_384PP.keys())[-1]
                    row_cells[2].text = str(volume)
                    volume_fulfilled = True

    # Assigning wells and volumes for level 1 plasmid backbone E (level 1)
    if int(GUI.transcription_unit_quantity_combo.get()) > 4:
        volume_fulfilled = False
        previous_fulfilment = 0
        while not volume_fulfilled:
            required_transfer_volume = 0
            for reaction in selected_reaction_ratios_level_1:
                if reaction == " reaction 1:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_1") * len(
                        MoClo.transcription_unit_5_names)
                elif reaction == " reaction 1:2":
                    required_transfer_volume += transfer_volume("level_1_backbone_1_2") * len(
                        MoClo.transcription_unit_5_names)
                elif reaction == " reaction 2:1":
                    required_transfer_volume += transfer_volume("level_1_backbone_2_1") * len(
                        MoClo.transcription_unit_5_names)
            row_cells = level_1_protocol_table.add_row().cells
            volume_required = required_transfer_volume - previous_fulfilment
            row_cells[1].text = "pTU1-E-RFP"
            if dead_volume + volume_required > 65000:
                volume = 65000
                assign_well_384(level_1_384PP, "pTU1-E-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                previous_fulfilment += 50000
                continue
            else:
                volume = dead_volume + volume_required
                assign_well_384(level_1_384PP, "pTU1-E-RFP", volume)
                row_cells[0].text = list(level_1_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                volume_fulfilled = True

    #################### level 1 6 RES source plate for deionised water ###########################
    level_1_prep_6res = document.add_paragraph("")
    level_1_prep_6res.add_run("\n" + "b) Add deionised water to the corresponding well(s) in the Echo®" +
                              " Qualified reservoir (6RES) as specified in the table below:")
    level_1_6res_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_1_6res_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "Reagent"
    row_1_cells[2].text = "Quantity (nl)"

    # Assigning well and volume for deionised water (level 1)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 250000
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += transfer_volume("level_1_water_1_1") * level_1_tu_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += transfer_volume("level_1_water_1_2") * level_1_tu_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += transfer_volume("level_1_water_2_1") * level_1_tu_quantity
        row_cells = level_1_6res_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "deionised water"
        if dead_volume + volume_required > 2800000:
            volume = 2800000
            assign_well_6res(level_1_6RES, "deionised water", volume)
            row_cells[0].text = list(level_1_6RES.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 2550000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_6res(level_1_6RES, "deionised water", volume)
            row_cells[0].text = list(level_1_6RES.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    ################# level 1 LDV source plate for other reagents (enzymes and buffers) #################
    level_1_prep_ldv = document.add_paragraph("")
    level_1_prep_ldv.add_run("\n" + "c) Add reagents to their corresponding wells in the" +
                             " Low Dead Volume Echo® Qualified 384-Well COC Source Microplate (384LDV) as" +
                             " specified in the table below:")
    level_1_ldv_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_1_ldv_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "Reagent"
    row_1_cells[2].text = "Quantity (nl)"

    # Assigning wells and volumes for DNA ligase buffer (level 1)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 3000
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += 500 * level_1_tu_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += 500 * level_1_tu_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += 500 * level_1_tu_quantity
        row_cells = level_1_ldv_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "10x DNA ligase buffer (Promega)"
        if dead_volume + volume_required > 12000:
            volume = 12000
            assign_well_384(level_1_LDV, "10x DNA ligase buffer (Promega)", volume)
            row_cells[0].text = list(level_1_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 9000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_384(level_1_LDV, "10x DNA ligase buffer (Promega)", volume)
            row_cells[0].text = list(level_1_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    # Assigning wells and volumes for DNA ligase (level 1)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 6000
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += 125 * level_1_tu_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += 125 * level_1_tu_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += 125 * level_1_tu_quantity
        row_cells = level_1_ldv_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "1-3 units T4 DNA ligase (Promega)"
        if dead_volume + volume_required > 14000:
            volume = 14000
            assign_well_384(level_1_LDV, "1-3 units T4 DNA ligase (Promega)", volume)
            row_cells[0].text = list(level_1_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 8000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_384(level_1_LDV, "1-3 units T4 DNA ligase (Promega)", volume)
            row_cells[0].text = list(level_1_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    # Assigning wells and volumes for BsaI-HF (level 1)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 6000
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += 250 * level_1_tu_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += 250 * level_1_tu_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += 250 * level_1_tu_quantity
        row_cells = level_1_ldv_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "BsaI-HF (NEB)"
        if dead_volume + volume_required > 14000:
            volume = 14000
            assign_well_384(level_1_LDV, "BsaI-HF (NEB)", volume)
            row_cells[0].text = list(level_1_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 8000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_384(level_1_LDV, "BsaI-HF (NEB)", volume)
            row_cells[0].text = list(level_1_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    ######################### Assigning wells in level 1 output plate ###############################
    level_1_output_intro = document.add_paragraph("")
    level_1_output_intro.add_run("\n" + "Locations of produced transcription unit variants in" +
                                 " output plate:").bold = True

    level_1_output_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_1_output_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "contained TU variant"
    row_1_cells[2].text = "volume (nl)"

    # Level 1 transcription unit variant 1
    for reaction in selected_reaction_ratios_level_1:
        for variant in MoClo.transcription_unit_1_names:
            row_cells = level_1_output_table.add_row().cells
            assign_well_384(level_1_output, variant + reaction, 5000)
            row_cells[0].text = list(level_1_output.keys())[-1]
            row_cells[1].text = variant + reaction
            row_cells[2].text = "5000"
    # Level 1 transcription unit variant 2
    for reaction in selected_reaction_ratios_level_1:
        for variant in MoClo.transcription_unit_2_names:
            row_cells = level_1_output_table.add_row().cells
            assign_well_384(level_1_output, variant + reaction, 5000)
            row_cells[0].text = list(level_1_output.keys())[-1]
            row_cells[1].text = variant + reaction
            row_cells[2].text = "5000"

    # Level 1 transcription unit variant 3
    for reaction in selected_reaction_ratios_level_1:
        for variant in MoClo.transcription_unit_3_names:
            row_cells = level_1_output_table.add_row().cells
            assign_well_384(level_1_output, variant + reaction, 5000)
            row_cells[0].text = list(level_1_output.keys())[-1]
            row_cells[1].text = variant + reaction
            row_cells[2].text = "5000"

    # Level 1 transcription unit variant 4
    for reaction in selected_reaction_ratios_level_1:
        for variant in MoClo.transcription_unit_4_names:
            row_cells = level_1_output_table.add_row().cells
            assign_well_384(level_1_output, variant + reaction, 5000)
            row_cells[0].text = list(level_1_output.keys())[-1]
            row_cells[1].text = variant + reaction
            row_cells[2].text = "5000"

    # Level 1 transcription unit variant 5
    for reaction in selected_reaction_ratios_level_1:
        for variant in MoClo.transcription_unit_5_names:
            row_cells = level_1_output_table.add_row().cells
            assign_well_384(level_1_output, variant + reaction, 5000)
            row_cells[0].text = list(level_1_output.keys())[-1]
            row_cells[1].text = variant + reaction
            row_cells[2].text = "5000"

    level_1_scripts = document.add_paragraph("")
    level_1_scripts.add_run("d) Load input plates into the Echo 525, and carry out protocols with the following "
                            "scripts:")
    level_1_scripts.add_run("\n" + "-" + protocol_name + "_6res.csv")
    level_1_scripts.add_run("\n" + "-" + protocol_name + "_ldv.csv")
    level_1_scripts.add_run("\n" + "-" + protocol_name + "_dna")
    level_1_transform = document.add_paragraph("")
    level_1_transform.add_run("e) Run a PCR protocol for the contents of the output plate, consisting of:")
    level_1_transform.add_run("\n" + "15-30 cycles of:" + "\n" + "15-30 cycles of:" + "\n" + "-5 minutes at 37°C" +
                              "\n" + "-10 minutes at 16°C" + "\n" + "Followed by (only once):" + "\n" + "-5 minutes "
                                                                                                        "at 50°C" +
                              "\n" + "-5 minutes at 80°C ")
    level_1_transform.add_run("\n" + "f) Transform 5 uL of each output mixture into 50uL of chemically" +
                              " compotent Escherichia coli (E. coli) dH10a by heat shock transformation" +
                              " (One mixture per cell culture)")

    document.add_heading("Creating level 2 transcription units (TUs)", 2)
    tu_2_prep_1_intro = document.add_paragraph("")
    tu_2_prep_1_intro.add_run("Plates required:").bold = True
    tu_2_prep_1_intro.add_run("\n" + "-2x Echo® Qualified 384-Well Polypropylene Source Microplate (384PP) (One of "
                                     "these is used as an output plate and must remain empty)")
    tu_2_prep_1_intro.add_run("\n" + "-Echo® Qualified 384-Well COC Source Microplate, Low Dead Volume (384LDV)")
    tu_2_prep_1_intro.add_run("\n" + "-Echo® Qualified Reservoir (6RES)" + "\n")
    tu_2_prep_1_intro.add_run("a) Recover and re-suspend level 1 transcription units from transformed cells in" +
                              " previous step" + "\n")
    tu_2_prep_1_intro.add_run("b) Add level 1 transcription units and level 2 plasmid backbones to their specified" +
                              " wells in the Echo® Qualified 384-Well Polypropylene Source Microplate (384PP) as" +
                              " specified in the table below:")
    level_2_protocol_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_2_protocol_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "Genetic part"
    row_1_cells[2].text = "Quantity (nl)"

    ########################## level 2 384 source plate for level 1 TU's and plasmid backbones ########################
    # Assigning wells and volumes for level 1 transcription units (level 2)
    tu1_quantities = MoClo.tu1_quantities
    dead_volume = 15000
    for variant in tu1_quantities.keys():
        volume_fulfilled = False
        previous_fulfilment = 0
        while not volume_fulfilled:
            required_transfer_volume = 0
            for reaction in selected_reaction_ratios_level_2:
                if reaction == " reaction 1:1":
                    required_transfer_volume += transfer_volume("level_2_tu_1_1") * tu1_quantities[variant]
                elif reaction == " reaction 1:2":
                    required_transfer_volume += transfer_volume("level_2_tu_1_2") * tu1_quantities[variant]
                elif reaction == " reaction 2:1":
                    required_transfer_volume += transfer_volume("level_2_tu_2_1") * tu1_quantities[variant]
            row_cells = level_2_protocol_table.add_row().cells
            volume_required = required_transfer_volume - previous_fulfilment
            row_cells[1].text = variant
            if dead_volume + volume_required > 65000:
                volume = 65000
                assign_well_384(level_2_384PP, variant, volume)
                row_cells[0].text = list(level_2_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                previous_fulfilment += 50000
                continue
            else:
                volume = dead_volume + volume_required
                assign_well_384(level_2_384PP, variant, volume)
                row_cells[0].text = list(level_2_384PP.keys())[-1]
                row_cells[2].text = str(volume)
                volume_fulfilled = True

    # Assigning wells and volumes for level 2 plasmid backbones (level 2)
    if int(GUI.transcription_unit_quantity_combo.get()) == 2:
        level_2_backbone = "pTU2-a-RFP"
    elif int(GUI.transcription_unit_quantity_combo.get()) == 3:
        level_2_backbone = "pTU2-b-RFP"
    elif int(GUI.transcription_unit_quantity_combo.get()) > 3:
        level_2_backbone = "pTU2-A-RFP"

    tu2_quantity = len(MoClo.level_2_names)
    dead_volume = 15000
    volume_fulfilled = False
    previous_fulfilment = 0
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_2:
            if reaction == " reaction 1:1":
                required_transfer_volume += transfer_volume("level_2_backbone_1_1") * tu2_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += transfer_volume("level_2_backbone_1_2") * tu2_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += transfer_volume("level_2_backbone_2_1") * tu2_quantity

        row_cells = level_2_protocol_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = level_2_backbone
        if dead_volume + volume_required > 65000:
            volume = 65000
            assign_well_384(level_2_384PP, level_2_backbone, volume)
            row_cells[0].text = list(level_2_384PP.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 50000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_384(level_2_384PP, level_2_backbone, volume)
            row_cells[0].text = list(level_2_384PP.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    ########################## level 2 6RES source plate for deionised water ########################
    level_2_prep_6res = document.add_paragraph("")
    level_2_prep_6res.add_run("\n" + "b) Add deionised water to the corresponding well(s) in the Echo®" +
                              " Qualified reservoir (6RES) as specified in the table below:")
    level_2_6res_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_2_6res_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "Reagent"
    row_1_cells[2].text = "Quantity (nl)"

    # Assigning wells and volumes for deionised water (level 1)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 250000
    single_transfer_volume = transfer_volume("level_2_water_2_1")
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += transfer_volume("level_2_water_1_1") * tu2_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += transfer_volume("level_2_water_1_2") * tu2_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += transfer_volume("level_2_water_2_1") * tu2_quantity
        row_cells = level_2_6res_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "deionised water"
        if dead_volume + volume_required > 2800000:
            volume = 2800000
            assign_well_6res(level_2_6RES, "deionised water", volume)
            row_cells[0].text = list(level_2_6RES.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 2550000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_6res(level_2_6RES, "deionised water", volume)
            row_cells[0].text = list(level_2_6RES.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    ################# level 2 LDV source plate for other reagents (enzymes and buffers) #################
    level_2_prep_ldv = document.add_paragraph("")
    level_2_prep_ldv.add_run("\n" + "c) Add reagents to their corresponding wells in the" +
                             " Low Dead Volume Echo® Qualified 384-Well COC Source Microplate (384LDV) as" +
                             " specified in the table below:")
    level_2_ldv_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_2_ldv_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "Reagent"
    row_1_cells[2].text = "Quantity (nl)"

    # Assigning wells and volumes for DNA ligase buffer (level 1)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 3000
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += 500 * tu2_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += 500 * tu2_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += 500 * tu2_quantity
        row_cells = level_2_ldv_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "10x DNA ligase buffer (Promega)"
        if dead_volume + volume_required > 12000:
            volume = 12000
            assign_well_384(level_2_LDV, "10x DNA ligase buffer (Promega)", volume)
            row_cells[0].text = list(level_2_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 9000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_384(level_2_LDV, "10x DNA ligase buffer (Promega)", volume)
            row_cells[0].text = list(level_2_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    # Assigning wells and volumes for DNA ligase (level 2)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 6000
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += 125 * tu2_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += 125 * tu2_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += 125 * tu2_quantity
        row_cells = level_2_ldv_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "1-3 units T4 DNA ligase (Promega)"
        if dead_volume + volume_required > 14000:
            volume = 14000
            assign_well_384(level_2_LDV, "1-3 units T4 DNA ligase (Promega)", volume)
            row_cells[0].text = list(level_2_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 8000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_384(level_2_LDV, "1-3 units T4 DNA ligase (Promega)", volume)
            row_cells[0].text = list(level_2_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    # Assigning wells and volumes for BsmBI (NEB) (level 2)
    volume_fulfilled = False
    previous_fulfilment = 0
    dead_volume = 6000
    while not volume_fulfilled:
        required_transfer_volume = 0
        for reaction in selected_reaction_ratios_level_1:
            if reaction == " reaction 1:1":
                required_transfer_volume += 250 * tu2_quantity
            elif reaction == " reaction 1:2":
                required_transfer_volume += 250 * tu2_quantity
            elif reaction == " reaction 2:1":
                required_transfer_volume += 250 * tu2_quantity
        row_cells = level_2_ldv_table.add_row().cells
        volume_required = required_transfer_volume - previous_fulfilment
        row_cells[1].text = "BsmBI (NEB)"
        if dead_volume + volume_required > 14000:
            volume = 14000
            assign_well_384(level_2_LDV, "BsmBI (NEB)", volume)
            row_cells[0].text = list(level_2_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            previous_fulfilment += 8000
            continue
        else:
            volume = dead_volume + volume_required
            assign_well_384(level_2_LDV, "BsmBI (NEB)", volume)
            row_cells[0].text = list(level_2_LDV.keys())[-1]
            row_cells[2].text = str(volume)
            volume_fulfilled = True

    ######################### Assigning wells in level 2 output plate ###############################
    level_2_output_intro = document.add_paragraph("")
    level_2_output_intro.add_run("\n" + "Locations of produced transcription unit variants in" +
                                 " output plate:").bold = True

    level_2_output_table = document.add_table(rows=1, cols=3)
    row_1_cells = level_2_output_table.rows[0].cells
    row_1_cells[0].text = "Well"
    row_1_cells[1].text = "contained TU variant"
    row_1_cells[2].text = "volume (nl)"

    # Level 2 transcription units
    for reaction in selected_reaction_ratios_level_2:
        for variant in MoClo.level_2_names:
            row_cells = level_2_output_table.add_row().cells
            assign_well_384(level_2_output, variant + reaction, 5000)
            row_cells[0].text = list(level_2_output.keys())[-1]
            row_cells[1].text = variant + reaction
            row_cells[2].text = "5000"

    level_2_scripts = document.add_paragraph("")
    level_2_scripts.add_run("d) Load input plates into the Echo 525, and carry out protocols with the following "
                            "scripts:")
    level_2_scripts.add_run("\n" + "-" + protocol_name + "_6res_2.csv")
    level_2_scripts.add_run("\n" + "-" + protocol_name + "_ldv_2.csv")
    level_2_scripts.add_run("\n" + "-" + protocol_name + "_dna_2")
    level_2_transform = document.add_paragraph("")
    level_2_transform.add_run("e) Incubate the contents of the output plate at 37°C for 16 hours")
    level_2_transform.add_run("\n" + "f) Transform 5 uL of each output mixture into 50uL of chemically" +
                              " compotent Escherichia coli (E. coli) dH10a by heat shock transformation" +
                              " (One mixture per cell culture)")


# Appendix of document, containing all parts, transcription units, and final designs
def create_appendix():
    # Parts
    document.add_page_break()
    document.add_heading("Appendix", 1)
    document.add_heading("Promoters", 2)
    counter = 0
    for promoter in MoClo.promoter_identities:
        document.add_heading(promoter, 3)
        description = document.add_paragraph("")
        description.add_run("Description: ").bold = True
        description.add_run(MoClo.promoter_descriptions[counter])
        backbone = document.add_paragraph("")
        backbone.add_run("Destination level 0 backbone: ").bold = True
        backbone.add_run("pBP-lacZ, available from 'http://www.addgene.org/72948/'")
        modifications = document.add_paragraph("")
        modifications.add_run("Notes:").bold = True
        for modification in MoClo.promoter_modifications[counter]:
            modifications.add_run("\n" + "-" + modification)
        design = document.add_paragraph("")
        design.add_run("Part design: ").bold = True
        if GUI.include_fusion_site_combo.get() == "Yes":
            design.add_run("\n" + "5' " + MoClo.promoter_sequences_1[counter] + " 3'" + "\n" +
                           "3'     " + MoClo.promoter_sequences_2[counter] + "         5'")
        if GUI.include_fusion_site_combo.get() == "No":
            design.add_run("\n" + "5' " + MoClo.promoter_sequences_1[counter] + " 3'" + "\n" +
                           "3' " + MoClo.promoter_sequences_2[counter] + " 5'")
        counter = counter + 1

    document.add_heading("Ribosome binding sites (RBSs)", 2)
    counter = 0
    for rbs in MoClo.rbs_identities:
        document.add_heading(rbs, 3)
        description = document.add_paragraph("")
        description.add_run("Description: ").bold = True
        description.add_run(MoClo.rbs_descriptions[counter])
        backbone = document.add_paragraph("")
        backbone.add_run("Destination level 0 backbone: ").bold = True
        backbone.add_run("pBP-lacZ, available from 'http://www.addgene.org/72948/'")
        modifications = document.add_paragraph("")
        modifications.add_run("Notes:").bold = True
        for modification in MoClo.rbs_modifications[counter]:
            modifications.add_run("\n" + "-" + modification)
        design = document.add_paragraph("")
        design.add_run("Part design: ").bold = True
        if GUI.include_fusion_site_combo.get() == "Yes":
            design.add_run("\n" + "5' " + MoClo.rbs_sequences_1[counter] + " 3'" + "\n" +
                           "3'     " + MoClo.rbs_sequences_2[counter] + "         5'")
        if GUI.include_fusion_site_combo.get() == "No":
            design.add_run("\n" + "5' " + MoClo.rbs_sequences_1[counter] + " 3'" + "\n" +
                           "3' " + MoClo.rbs_sequences_2[counter] + " 5'")
        counter = counter + 1

    if GUI.include_signal_combo.get() == "Yes":
        document.add_heading("Signal peptides", 2)
        counter = 0
        for signal in MoClo.signal_identities:
            document.add_heading(signal, 3)
            description = document.add_paragraph("")
            description.add_run("Description: ").bold = True
            description.add_run(MoClo.signal_descriptions[counter])
            backbone = document.add_paragraph("")
            backbone.add_run("Destination level 0 backbone: ").bold = True
            backbone.add_run("pBP-lacZ, available from 'http://www.addgene.org/72948/'")
            modifications = document.add_paragraph("")
            modifications.add_run("Notes:").bold = True
            for modification in MoClo.signal_modifications[counter]:
                modifications.add_run("\n" + "-" + modification)
            design = document.add_paragraph("")
            design.add_run("Part design: ").bold = True
            if GUI.include_fusion_site_combo.get() == "Yes":
                design.add_run("\n" + "5' " + MoClo.signal_sequences_1[counter] + " 3'" + "\n" +
                               "3'     " + MoClo.signal_sequences_2[counter] + "         5'")
            if GUI.include_fusion_site_combo.get() == "No":
                design.add_run("\n" + "5' " + MoClo.signal_sequences_1[counter] + " 3'" + "\n" +
                               "3' " + MoClo.signal_sequences_2[counter] + " 5'")
            counter = counter + 1

    document.add_heading("Coding regions (CDSs)", 2)
    counter = 0
    for cds in MoClo.cds_identities:
        document.add_heading(cds, 3)
        description = document.add_paragraph("")
        description.add_run("Description: ").bold = True
        description.add_run(MoClo.cds_descriptions[counter])
        backbone = document.add_paragraph("")
        backbone.add_run("Destination level 0 backbone: ").bold = True
        backbone.add_run("pBP-ORF, available from 'http://www.addgene.org/72949/'")
        modifications = document.add_paragraph("")
        modifications.add_run("Notes:").bold = True
        for modification in MoClo.cds_modifications[counter]:
            modifications.add_run("\n" + "-" + modification)
        design = document.add_paragraph("")
        design.add_run("Part design: ").bold = True
        if GUI.include_fusion_site_combo.get() == "Yes":
            design.add_run("\n" + "5' " + MoClo.cds_sequences_1[counter] + " 3'" + "\n" +
                           "3'     " + MoClo.cds_sequences_2[counter] + "         5'")
        if GUI.include_fusion_site_combo.get() == "No":
            design.add_run("\n" + "5' " + MoClo.cds_sequences_1[counter] + " 3'" + "\n" +
                           "3' " + MoClo.cds_sequences_2[counter] + " 5'")
        counter = counter + 1

    document.add_heading("Terminators", 2)
    counter = 0
    for terminator in MoClo.terminator_identities:
        document.add_heading(terminator, 3)
        description = document.add_paragraph("")
        description.add_run("Description: ").bold = True
        description.add_run(MoClo.terminator_descriptions[counter])
        backbone = document.add_paragraph("")
        backbone.add_run("Destination level 0 backbone: ").bold = True
        backbone.add_run("pBP-lacZ, available from 'http://www.addgene.org/72948/'")
        modifications = document.add_paragraph("")
        modifications.add_run("Notes:").bold = True
        for modification in MoClo.terminator_modifications[counter]:
            modifications.add_run("\n" + "-" + modification)
        design = document.add_paragraph("")
        design.add_run("Part design: ").bold = True
        if GUI.include_fusion_site_combo.get() == "Yes":
            design.add_run("\n" + "5' " + MoClo.terminator_sequences_1[counter] + " 3'" + "\n" +
                           "3'     " + MoClo.terminator_sequences_2[counter] + "         5'")
        if GUI.include_fusion_site_combo.get() == "No":
            design.add_run("\n" + "5' " + MoClo.terminator_sequences_1[counter] + " 3'" + "\n" +
                           "3' " + MoClo.terminator_sequences_2[counter] + " 5'")
        counter = counter + 1

    # Level 1 transcription units
    document.add_heading("Level 1 transcription units", 2)
    if int(GUI.transcription_unit_quantity_combo.get()) > 1:
        counter = 0
        for variant in MoClo.transcription_unit_1_names:
            document.add_heading(variant, 3)
            part_list = document.add_paragraph("")
            part_list.add_run("Parts: ").bold = True
            for part in MoClo.transcription_unit_1_part_id[counter]:
                part_list.add_run(part)
                part_list.add_run(", ")
            backbone = document.add_paragraph("")
            backbone.add_run("Plasmid backbone: ").bold = True
            backbone.add_run("pTU1-A-RFP, available from 'http://www.addgene.org/72939/'")
            notes = document.add_paragraph("")
            notes.add_run("Notes: ").bold = True
            notes.add_run(MoClo.transcription_unit_1_notes[counter])
            sequences = document.add_paragraph("")
            sequences.add_run("Sequence (Excluding plasmid backbone): ").bold = True
            sequences.add_run("5' ")
            sequences.add_run(MoClo.transcription_unit_1_sequences[counter])
            sequences.add_run(" 3'")
            counter = counter + 1

        counter = 0
        for variant in MoClo.transcription_unit_2_names:
            document.add_heading(variant, 3)
            part_list = document.add_paragraph("")
            part_list.add_run("Parts: ").bold = True
            for part in MoClo.transcription_unit_2_part_id[counter]:
                part_list.add_run(part)
                part_list.add_run(", ")
            backbone = document.add_paragraph("")
            backbone.add_run("Plasmid backbone: ").bold = True
            backbone.add_run("pTU1-B-RFP, available from 'http://www.addgene.org/72940/'")
            notes = document.add_paragraph("")
            notes.add_run("Notes: ").bold = True
            notes.add_run(MoClo.transcription_unit_2_notes[counter])
            sequences = document.add_paragraph("")
            sequences.add_run("Sequence (Excluding plasmid backbone): ").bold = True
            sequences.add_run("5' ")
            sequences.add_run(MoClo.transcription_unit_2_sequences[counter])
            sequences.add_run(" 3'")
            counter = counter + 1

    if int(GUI.transcription_unit_quantity_combo.get()) > 2:
        counter = 0
        for variant in MoClo.transcription_unit_3_names:
            document.add_heading(variant, 3)
            part_list = document.add_paragraph("")
            part_list.add_run("Parts: ").bold = True
            for part in MoClo.transcription_unit_3_part_id[counter]:
                part_list.add_run(part)
                part_list.add_run(", ")
            backbone = document.add_paragraph("")
            backbone.add_run("Plasmid backbone: ").bold = True
            backbone.add_run("pTU1-C-RFP, available from 'http://www.addgene.org/72941/'")
            notes = document.add_paragraph("")
            notes.add_run("Notes: ").bold = True
            notes.add_run(MoClo.transcription_unit_3_notes[counter])
            sequences = document.add_paragraph("")
            sequences.add_run("Sequence (Excluding plasmid backbone): ").bold = True
            sequences.add_run("5' ")
            sequences.add_run(MoClo.transcription_unit_3_sequences[counter])
            sequences.add_run(" 3'")
            counter = counter + 1

    if int(GUI.transcription_unit_quantity_combo.get()) > 3:
        counter = 0
        for variant in MoClo.transcription_unit_4_names:
            document.add_heading(variant, 3)
            part_list = document.add_paragraph("")
            part_list.add_run("Parts: ").bold = True
            for part in MoClo.transcription_unit_4_part_id[counter]:
                part_list.add_run(part)
                part_list.add_run(", ")
            backbone = document.add_paragraph("")
            backbone.add_run("Plasmid backbone: ").bold = True
            if GUI.transcription_unit_quantity_combo.get() == "4":
                backbone.add_run("pTU1-D-RFP, available from 'http://www.addgene.org/72942/'")
            if GUI.transcription_unit_quantity_combo.get() == "5":
                backbone.add_run("pTU1-D1-RFP, available from 'http://www.addgene.org/72943/'")
            notes = document.add_paragraph("")
            notes.add_run("Notes: ").bold = True
            notes.add_run(MoClo.transcription_unit_4_notes[counter])
            sequences = document.add_paragraph("")
            sequences.add_run("Sequence (Excluding plasmid backbone): ").bold = True
            sequences.add_run("5' ")
            sequences.add_run(MoClo.transcription_unit_4_sequences[counter])
            sequences.add_run(" 3'")
            counter = counter + 1

    if int(GUI.transcription_unit_quantity_combo.get()) > 4:
        counter = 0
        for variant in MoClo.transcription_unit_5_names:
            document.add_heading(variant, 3)
            part_list = document.add_paragraph("")
            part_list.add_run("Parts: ").bold = True
            for part in MoClo.transcription_unit_5_part_id[counter]:
                part_list.add_run(part)
                part_list.add_run(", ")
            backbone = document.add_paragraph("")
            backbone.add_run("Plasmid backbone: ").bold = True
            backbone.add_run("pTU1-E-RFP, available from 'http://www.addgene.org/72944/'")
            notes = document.add_paragraph("")
            notes.add_run("Notes: ").bold = True
            notes.add_run(MoClo.transcription_unit_5_notes[counter])
            sequences = document.add_paragraph("")
            sequences.add_run("Sequence (Excluding plasmid backbone): ").bold = True
            sequences.add_run("5' ")
            sequences.add_run(MoClo.transcription_unit_5_sequences[counter])
            sequences.add_run(" 3'")
            counter = counter + 1

    # Level 2 transcription units
    counter = 0
    document.add_heading("Level 2 multi-TU constructs", 2)
    selected_vector = document.add_paragraph("")
    selected_vector.add_run("Level 2 plasmid backbone: ").bold = True
    selected_vector.add_run(MoClo.level_2_vector_name)
    for variant in MoClo.level_2_names:
        document.add_heading(variant, 3)
        sub_units = document.add_paragraph("")
        sub_units.add_run("Sub-units: ").bold = True
        for unit in MoClo.level_2_sub_units[counter]:
            sub_units.add_run(unit)
            sub_units.add_run(", ")
        sequences = document.add_paragraph("")
        sequences.add_run("Sequence (Excluding plasmid backbone): ").bold = True
        sequences.add_run("5' ")
        sequences.add_run(MoClo.level_2_sequences[counter])
        sequences.add_run(" 3'")
        counter = counter + 1


# Create EcoFlex protocol
def create_protocol(event):
    get_variables()
    MoClo.swap_codons_ecoflex()
    MoClo.check_biopart_sites_ecoflex()
    MoClo.ecoflex_fusion_sites()
    MoClo.create_transcription_unit_variants()
    MoClo.final_oligonucleotides_2()
    MoClo.transcription_unit_format()
    MoClo.part_use_quantity()
    MoClo.level_2_format()
    MoClo.transcription_unit_use_quantity()
    calculate_part_quantity()
    calculate_level_1_quantity()
    create_selected_reaction_list()
    title_introduction()
    if GUI.assembly_method_combo.get() == "Manual":
        create_manual_protocol()
    else:
        create_automatic_protocol()
        from EcoFlex_scripts import create_scripts
        create_scripts()

    create_appendix()
    document.save("protocols_and_scripts\\" + protocol_name + ".docx")
    GUI.moclo_reset("<Button-1>")
    clear_protocol_globals()
    GUI.successful_creation()


# Clear ecoflex protocol generation globals
def clear_protocol_globals():
    global document
    document = Document()
    global level_1_384PP
    level_1_384PP = {}
    global level_1_LDV
    level_1_LDV = {}
    global level_1_6RES
    level_1_6RES = {}
    global level_1_output
    level_1_output = {}
    global level_2_384PP
    level_2_384PP = {}
    global level_2_LDV
    level_2_LDV = {}
    global level_2_6RES
    level_2_6RES = {}
    global level_2_output
    level_2_output = {}
